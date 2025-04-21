import os
os.environ['TORCH_DISABLE_CUSTOM_CLASS_PATHS'] = '1'

# Standard imports
import hashlib
from typing import Dict, List, Optional, Union
from pathlib import Path
from diskcache import Cache
import json
import pandas as pd
from PyPDF2 import PdfReader
from docx import Document
import spacy
from rapidfuzz import fuzz

# Lazy load PyTorch-dependent imports
sentence_transformer = None
def get_sentence_transformer():
    global sentence_transformer
    if sentence_transformer is None:
        from sentence_transformers import SentenceTransformer
        sentence_transformer = SentenceTransformer('all-MiniLM-L6-v2')
    return sentence_transformer

from sklearn.metrics.pairwise import cosine_similarity
import numpy as np
from langchain_community.document_loaders import (
    PyPDFLoader,
    CSVLoader,
    Docx2txtLoader,
    TextLoader
)
from langchain.text_splitter import RecursiveCharacterTextSplitter
from config import (
    CACHE_DIR,
    TEXT_CHUNK_SIZE,
    TEXT_CHUNK_OVERLAP,
    SUPPORTED_ENCODINGS
)

class DocumentProcessor:
    def __init__(self, cache_dir: str = CACHE_DIR):
        self.cache = Cache(cache_dir)
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=TEXT_CHUNK_SIZE,
            chunk_overlap=TEXT_CHUNK_OVERLAP,
            separators=["\n\n", "\n", ".", " ", ""]
        )
        
        # Initialize NLP components
        try:
            self.nlp = spacy.load("en_core_web_sm")
        except:
            os.system("python -m spacy download en_core_web_sm")
            self.nlp = spacy.load("en_core_web_sm")
        
        # Don't initialize sentence transformer immediately
        self.sentence_model = None
        
        # Define requirement indicators
        self.requirement_indicators = {
            'mandatory': ['shall', 'must', 'required', 'mandatory'],
            'recommended': ['should', 'recommended', 'may'],
            'measurement': ['measure', 'assess', 'evaluate', 'monitor'],
            'validation': ['verify', 'validate', 'test', 'audit']
        }

    def _get_sentence_model(self):
        if self.sentence_model is None:
            self.sentence_model = get_sentence_transformer()
        return self.sentence_model

    def process_document(self, file_path: Union[str, Path]) -> Dict[str, any]:
        """Process a document with enhanced requirement extraction"""
        file_path = str(file_path)
        file_hash = self._calculate_file_hash(file_path)
        
        # Try to get from cache first
        cached_result = self.cache.get(file_hash)
        if cached_result is not None:
            return cached_result
            
        # Process the file based on its type
        result = self._process_file(file_path)
        
        # Cache the result
        self.cache.set(file_hash, result)
        return result
    
    def _calculate_file_hash(self, file_path: str) -> str:
        """Calculate a hash of the file content for caching"""
        hasher = hashlib.sha256()
        with open(file_path, 'rb') as f:
            while chunk := f.read(8192):
                hasher.update(chunk)
        return hasher.hexdigest()
    
    def _process_file(self, file_path: str) -> Dict[str, any]:
        """Process different file types with enhanced text analysis"""
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == '.csv':
            return self._process_csv(file_path)
        elif ext == '.pdf':
            return self._process_pdf(file_path)
        elif ext in ['.doc', '.docx']:
            return self._process_word(file_path)
        elif ext == '.txt':
            return self._process_text(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")
    
    def _process_csv(self, file_path: str) -> Dict[str, any]:
        """Process CSV files"""
        df = pd.read_csv(file_path)
        return {
            "type": "csv",
            "content": df.to_dict('records'),
            "metadata": {
                "columns": list(df.columns),
                "row_count": len(df)
            }
        }
    
    def _process_pdf(self, file_path: str) -> Dict[str, any]:
        """Process PDF files"""
        reader = PdfReader(file_path)
        content = []
        metadata = {}
        
        # Extract text and metadata from each page
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            content.append({
                "page": i + 1,
                "text": text,
                "layout": self._analyze_page_layout(text)
            })
        
        if reader.metadata:
            metadata = {k.lower(): v for k, v in reader.metadata.items()}
        
        return {
            "type": "pdf",
            "content": content,
            "metadata": metadata
        }
    
    def _process_word(self, file_path: str) -> Dict[str, any]:
        """Process Word documents"""
        doc = Document(file_path)
        content = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                content.append({
                    "text": para.text,
                    "style": para.style.name
                })
        
        return {
            "type": "docx",
            "content": content,
            "metadata": {
                "sections": len(doc.sections),
                "paragraphs": len(doc.paragraphs)
            }
        }
    
    def _process_text(self, file_path: str) -> Dict[str, any]:
        """Process text file and extract controls"""
        loader = TextLoader(file_path)
        document = loader.load()
        chunks = self.text_splitter.split_documents(document)
        
        controls = []
        for chunk in chunks:
            controls.extend(self._parse_control_section(chunk.page_content))
        return {
            "type": "txt",
            "content": controls
        }
    
    def _process_text_content(self, text: str) -> List[Dict[str, any]]:
        """Process text content with advanced NLP analysis"""
        doc = self.nlp(text)
        requirements = []
        
        # Split into meaningful chunks
        chunks = [sent.text.strip() for sent in doc.sents]
        
        for chunk in chunks:
            # Skip empty or very short chunks
            if len(chunk.split()) < 5:
                continue
                
            # Analyze requirement characteristics
            req_info = self._analyze_requirement(chunk)
            
            if req_info['is_requirement']:
                requirements.append({
                    'content': chunk,
                    'type': req_info['type'],
                    'confidence': req_info['confidence'],
                    'indicators': req_info['indicators'],
                    'entities': req_info['entities'],
                    'ambiguity_score': req_info['ambiguity_score']
                })
        
        return requirements
    
    def _analyze_requirement(self, text: str) -> Dict[str, any]:
        """Analyze text for requirement characteristics"""
        doc = self.nlp(text.lower())
        
        # Check for requirement indicators
        indicators = []
        for category, terms in self.requirement_indicators.items():
            found_terms = [term for term in terms if term in text.lower()]
            if found_terms:
                indicators.append({
                    'category': category,
                    'terms': found_terms
                })
        
        # Extract named entities
        entities = [{
            'text': ent.text,
            'label': ent.label_,
            'start': ent.start_char,
            'end': ent.end_char
        } for ent in doc.ents]
        
        # Calculate ambiguity score
        ambiguous_terms = [
            'appropriate', 'adequate', 'sufficient', 'reasonable',
            'etc', 'and/or', 'as needed', 'if possible'
        ]
        ambiguity_score = sum(1 for term in ambiguous_terms if term in text.lower()) / len(ambiguous_terms)
        
        # Determine if text is likely a requirement
        is_requirement = bool(indicators) or any(self._is_requirement_verb(token) for token in doc)
        
        # Calculate confidence score
        confidence = self._calculate_requirement_confidence(text, indicators, ambiguity_score)
        
        # Determine requirement type
        req_type = self._determine_requirement_type(indicators)
        
        return {
            'is_requirement': is_requirement,
            'type': req_type,
            'confidence': confidence,
            'indicators': indicators,
            'entities': entities,
            'ambiguity_score': ambiguity_score
        }
    
    def _is_requirement_verb(self, token) -> bool:
        """Check if a verb typically indicates a requirement"""
        requirement_verbs = {
            'implement', 'establish', 'maintain', 'ensure',
            'define', 'document', 'review', 'monitor'
        }
        return token.pos_ == 'VERB' and token.lemma_.lower() in requirement_verbs
    
    def _calculate_requirement_confidence(self, text: str, indicators: List[Dict[str, any]], ambiguity_score: float) -> float:
        """Calculate confidence score for requirement identification"""
        scores = []
        
        # Score based on requirement indicators
        if indicators:
            scores.append(min(1.0, len(indicators) * 0.3))
        
        # Score based on sentence structure
        doc = self.nlp(text)
        has_subject = any(token.dep_ == 'nsubj' for token in doc)
        has_verb = any(token.pos_ == 'VERB' for token in doc)
        has_object = any(token.dep_ in ['dobj', 'pobj'] for token in doc)
        structure_score = (has_subject + has_verb + has_object) / 3
        scores.append(structure_score)
        
        # Score based on clarity (inverse of ambiguity)
        clarity_score = 1 - ambiguity_score
        scores.append(clarity_score)
        
        # Calculate weighted average
        weights = [0.4, 0.3, 0.3]  # Weights for indicators, structure, clarity
        return sum(score * weight for score, weight in zip(scores, weights))
    
    def _determine_requirement_type(self, indicators: List[Dict[str, any]]) -> str:
        """Determine the type of requirement based on indicators"""
        if not indicators:
            return 'implicit'
            
        # Count indicator categories
        category_counts = {}
        for indicator in indicators:
            category_counts[indicator['category']] = category_counts.get(indicator['category'], 0) + 1
        
        # Determine primary type
        if category_counts.get('mandatory', 0) > 0:
            return 'mandatory'
        elif category_counts.get('recommended', 0) > 0:
            return 'recommended'
        elif category_counts.get('measurement', 0) > 0:
            return 'measurement'
        elif category_counts.get('validation', 0) > 0:
            return 'validation'
        else:
            return 'informative'
    
    def _calculate_semantic_similarity(self, text1: str, text2: str) -> float:
        """Calculate semantic similarity between two texts"""
        model = self._get_sentence_model()
        # Generate embeddings
        emb1 = model.encode([text1])[0]
        emb2 = model.encode([text2])[0]
        
        # Calculate cosine similarity
        similarity = cosine_similarity([emb1], [emb2])[0][0]
        return float(similarity)
    
    def _analyze_page_layout(self, text: str) -> Dict[str, any]:
        """Analyze the layout structure of text"""
        lines = text.split('\n')
        return {
            "line_count": len(lines),
            "avg_line_length": sum(len(line) for line in lines) / len(lines) if lines else 0,
            "indentation_levels": self._detect_indentation_levels(lines)
        }
    
    def _detect_indentation_levels(self, lines: List[str]) -> Dict[int, int]:
        """Detect and count different indentation levels"""
        indentation_counts = {}
        for line in lines:
            leading_spaces = len(line) - len(line.lstrip())
            if leading_spaces > 0:
                indentation_counts[leading_spaces] = indentation_counts.get(leading_spaces, 0) + 1
        return indentation_counts

    def _parse_control_section(self, text: str) -> List[Dict[str, str]]:
        """Parse a section of text to extract control information"""
        controls = []
        lines = text.split('\n')
        current_control = None
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # Check if line starts with a control identifier
            if self._is_control_identifier(line):
                if current_control:
                    controls.append(current_control)
                current_control = {'identifier': line, 'content': ''}
            elif current_control:
                current_control['content'] += f" {line}"
        
        # Add the last control if exists
        if current_control:
            controls.append(current_control)
            
        return controls

    def _is_control_identifier(self, text: str) -> bool:
        """Check if text represents an ISO 27001 control identifier"""
        text = text.strip()
        # Match patterns like "5.1", "A.5.1.1", etc.
        return (
            bool(text) and
            (text[0].isdigit() or text.startswith('A.')) and
            any(char.isdigit() for char in text) and
            len(text) <= 10  # Reasonable length for a control ID
        )